"""Layer 10 - Export: DocxExporter.

Single responsibility: produce output.docx from already-assembled text.
Right-to-left paragraphs and runs so Urdu renders and edits naturally in
Word, not as a left-to-right jumble.

RTL-setting approach (bidi paragraph element, run.font.rtl, bidi language
tag) ported from the old app/core/exporters/docx_exporter.py rather than
re-derived - per PROJECT_SPEC.md Section 4/7, that logic was already
correct and tested; only the input shape changed (a plain assembled string
here, not a DocumentResult with Page/OCRWord objects).

Lines prefixed with text_exporter.HEADING_MARKER (set by
assemble_text_with_headings, based on heading_classifier's relative-height
classification) get Word's real "Heading 1" paragraph style instead of a
body paragraph - that's what lets Word's automatic Table of Contents
(Insert > Table of Contents) find them; bold/larger text alone would not be
picked up by that feature. The marker character itself is stripped and
never appears in the visible run text.

Words wrapped in text_exporter.LOW_CONFIDENCE_START/END (set by
assemble_text_with_headings, based on low_confidence_flagger's per-word
confidence check) get a highlighted run instead of the surrounding
paragraph's plain formatting - a visible flag per PROJECT_SPEC.md Section 2
rule 2 ("never guess ... flag it"), without altering the recognized text
itself. The bracket characters themselves are stripped and never appear in
the visible run text, same as the heading marker.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Pt

from app.core.export.text_exporter import HEADING_MARKER, LOW_CONFIDENCE_END, LOW_CONFIDENCE_START

_LOW_CONFIDENCE_PATTERN = re.compile(re.escape(LOW_CONFIDENCE_START) + r"(.*?)" + re.escape(LOW_CONFIDENCE_END))


def _set_paragraph_rtl(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)


def _split_low_confidence_segments(line: str) -> list[tuple[str, bool]]:
    """Splits a line into (text, is_low_confidence) segments, stripping the
    LOW_CONFIDENCE_START/END markers themselves out of the text."""
    segments: list[tuple[str, bool]] = []
    pos = 0
    for match in _LOW_CONFIDENCE_PATTERN.finditer(line):
        if match.start() > pos:
            segments.append((line[pos:match.start()], False))
        segments.append((match.group(1), True))
        pos = match.end()
    if pos < len(line):
        segments.append((line[pos:], False))
    return segments


class DocxExporter:
    def export(self, text: str, path: Path) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        document = Document()

        for line in text.split("\n"):
            if not line.strip():
                continue
            is_heading = line.startswith(HEADING_MARKER)
            if is_heading:
                line = line[len(HEADING_MARKER):]
            if not line.strip():
                continue

            paragraph = document.add_paragraph(style="Heading 1" if is_heading else None)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_paragraph_rtl(paragraph)

            for segment_text, is_low_confidence in _split_low_confidence_segments(line):
                if not segment_text:
                    continue
                run = paragraph.add_run(segment_text)
                run.font.rtl = True
                run.font.size = Pt(20) if is_heading else Pt(14)
                if is_heading:
                    run.font.bold = True
                if is_low_confidence:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                rpr = run._element.get_or_add_rPr()
                lang = OxmlElement("w:lang")
                lang.set(qn("w:bidi"), "ur-PK")
                rpr.append(lang)

        document.save(str(path))
