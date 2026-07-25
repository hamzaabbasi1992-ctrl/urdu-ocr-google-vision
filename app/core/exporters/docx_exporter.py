"""output.docx - literal OCR text as a Word document with right-to-left
paragraphs/runs so Urdu text renders and edits naturally in Word."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Pt

from app.core.models import DocumentResult


def _set_paragraph_rtl(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)


def export_docx(result: DocumentResult, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()

    for page in result.pages:
        heading = document.add_heading(f"Page {page.page_number}", level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for line in page.text.split("\n"):
            if not line.strip():
                continue
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_paragraph_rtl(paragraph)
            run = paragraph.add_run(line)
            run.font.rtl = True
            run.font.size = Pt(14)
            rpr = run._element.get_or_add_rPr()
            lang = OxmlElement("w:lang")
            lang.set(qn("w:bidi"), "ur-PK")
            rpr.append(lang)

        if page is not result.pages[-1]:
            document.add_page_break()

    document.save(str(output_path))
