"""Standalone test for Layer 10's DocxExporter - verifies it writes a
valid .docx with the right text content and RTL formatting applied."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from app.core.export.docx_exporter import DocxExporter
from app.core.export.text_exporter import HEADING_MARKER, LOW_CONFIDENCE_END, LOW_CONFIDENCE_START


def test_export_creates_readable_docx_with_correct_paragraph_text(tmp_path: Path) -> None:
    path = tmp_path / "output.docx"
    DocxExporter().export("بسم الله\nالحمد للہ", path)

    assert path.exists()
    document = Document(str(path))
    paragraph_texts = [p.text for p in document.paragraphs if p.text.strip()]
    assert paragraph_texts == ["بسم الله", "الحمد للہ"]


def test_export_skips_blank_lines(tmp_path: Path) -> None:
    path = tmp_path / "output.docx"
    DocxExporter().export("line one\n\nline two", path)

    document = Document(str(path))
    paragraph_texts = [p.text for p in document.paragraphs if p.text.strip()]
    assert paragraph_texts == ["line one", "line two"]


def test_export_sets_rtl_bidi_on_paragraphs(tmp_path: Path) -> None:
    from docx.oxml.ns import qn

    path = tmp_path / "output.docx"
    DocxExporter().export("متن", path)

    document = Document(str(path))
    paragraph = next(p for p in document.paragraphs if p.text.strip())
    pPr = paragraph._p.pPr
    assert pPr is not None
    bidi = pPr.find(qn("w:bidi"))
    assert bidi is not None, "expected a w:bidi element marking the paragraph as right-to-left"


def test_export_creates_parent_directories(tmp_path: Path) -> None:
    path = tmp_path / "nested" / "dir" / "output.docx"
    DocxExporter().export("text", path)
    assert path.exists()


def test_export_empty_text_produces_valid_empty_document(tmp_path: Path) -> None:
    path = tmp_path / "output.docx"
    DocxExporter().export("", path)

    assert path.exists()
    document = Document(str(path))
    assert all(not p.text.strip() for p in document.paragraphs)


def test_export_applies_heading_1_style_to_marked_lines(tmp_path: Path) -> None:
    path = tmp_path / "output.docx"
    DocxExporter().export(f"{HEADING_MARKER}Chapter One\nbody text", path)

    document = Document(str(path))
    paragraphs = [p for p in document.paragraphs if p.text.strip()]
    assert paragraphs[0].text == "Chapter One"
    assert paragraphs[0].style.name == "Heading 1"
    assert paragraphs[1].text == "body text"
    assert paragraphs[1].style.name != "Heading 1"


def test_export_heading_marker_never_appears_in_visible_text(tmp_path: Path) -> None:
    path = tmp_path / "output.docx"
    DocxExporter().export(f"{HEADING_MARKER}Title\n{HEADING_MARKER}Another Title", path)

    document = Document(str(path))
    for paragraph in document.paragraphs:
        assert HEADING_MARKER not in paragraph.text


def test_export_highlights_low_confidence_run(tmp_path: Path) -> None:
    path = tmp_path / "output.docx"
    DocxExporter().export(f"good {LOW_CONFIDENCE_START}bad{LOW_CONFIDENCE_END}", path)

    document = Document(str(path))
    paragraph = next(p for p in document.paragraphs if p.text.strip())
    assert paragraph.text == "good bad"
    runs = paragraph.runs
    highlighted = [r for r in runs if r.font.highlight_color == WD_COLOR_INDEX.YELLOW]
    not_highlighted = [r for r in runs if r.font.highlight_color != WD_COLOR_INDEX.YELLOW]
    assert "".join(r.text for r in highlighted) == "bad"
    assert "".join(r.text for r in not_highlighted) == "good "


def test_export_low_confidence_markers_never_appear_in_visible_text(tmp_path: Path) -> None:
    path = tmp_path / "output.docx"
    DocxExporter().export(f"{LOW_CONFIDENCE_START}bad{LOW_CONFIDENCE_END} good", path)

    document = Document(str(path))
    for paragraph in document.paragraphs:
        assert LOW_CONFIDENCE_START not in paragraph.text
        assert LOW_CONFIDENCE_END not in paragraph.text


def test_export_non_flagged_line_has_no_highlighted_runs(tmp_path: Path) -> None:
    path = tmp_path / "output.docx"
    DocxExporter().export("plain body text", path)

    document = Document(str(path))
    paragraph = next(p for p in document.paragraphs if p.text.strip())
    assert all(r.font.highlight_color != WD_COLOR_INDEX.YELLOW for r in paragraph.runs)


def test_export_heading_with_low_confidence_word_keeps_both_formattings(tmp_path: Path) -> None:
    path = tmp_path / "output.docx"
    DocxExporter().export(f"{HEADING_MARKER}Chapter {LOW_CONFIDENCE_START}One{LOW_CONFIDENCE_END}", path)

    document = Document(str(path))
    paragraph = next(p for p in document.paragraphs if p.text.strip())
    assert paragraph.text == "Chapter One"
    assert paragraph.style.name == "Heading 1"
    flagged_run = next(r for r in paragraph.runs if r.text == "One")
    assert flagged_run.font.highlight_color == WD_COLOR_INDEX.YELLOW
    assert flagged_run.font.bold is True
